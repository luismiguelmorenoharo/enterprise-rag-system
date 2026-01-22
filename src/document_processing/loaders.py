"""
Multi-format document loaders.
Handles PDF, DOCX, and TXT with robust error handling.
"""

import PyPDF2
import docx
from pathlib import Path
from typing import Dict, Optional
import logging

from src.utils.logger import get_logger

logger = get_logger(__name__)


class DocumentLoader:
    """
    Load and extract text from multiple document formats.
    
    Supports:
    - PDF (with metadata extraction)
    - DOCX (paragraphs + tables)
    - TXT (plain text)
    """
    
    @staticmethod
    def load_pdf(file_path: Path) -> Dict[str, any]:
        """
        Extract text and metadata from PDF.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Dict with 'text' and 'metadata' keys
            
        Raises:
            ValueError: If PDF is encrypted or corrupted
        """
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Check if encrypted
                if reader.is_encrypted:
                    raise ValueError(f"PDF is encrypted: {file_path.name}")
                
                # Extract text from all pages
                text_parts = []
                for page_num, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                    else:
                        logger.warning(f"Page {page_num} has no extractable text")
                
                text = "\n\n".join(text_parts)
                
                # Extract metadata
                metadata = {
                    'pages': len(reader.pages),
                    'source': str(file_path),
                    'filename': file_path.name,
                    'format': 'pdf'
                }
                
                # Add PDF metadata if available
                if reader.metadata:
                    metadata.update({
                        'author': reader.metadata.get('/Author', 'Unknown'),
                        'title': reader.metadata.get('/Title', file_path.stem),
                        'creator': reader.metadata.get('/Creator', 'Unknown')
                    })
                else:
                    metadata['title'] = file_path.stem
                
                logger.info(
                    f"✓ Loaded PDF: {file_path.name} "
                    f"({len(text)} chars, {metadata['pages']} pages)"
                )
                
                return {
                    'text': text.strip(),
                    'metadata': metadata
                }
                
        except PyPDF2.errors.PdfReadError as e:
            logger.error(f"Invalid PDF file {file_path}: {str(e)}")
            raise ValueError(f"Corrupted PDF: {file_path.name}")
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def load_docx(file_path: Path) -> Dict[str, any]:
        """
        Extract text from DOCX including tables.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dict with 'text' and 'metadata' keys
        """
        try:
            doc = docx.Document(file_path)
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                
                if table_text:
                    text_parts.append("\n".join(table_text))
            
            text = "\n\n".join(text_parts)
            
            # Metadata
            metadata = {
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'source': str(file_path),
                'filename': file_path.name,
                'format': 'docx',
                'title': file_path.stem
            }
            
            # Try to get core properties
            try:
                core_props = doc.core_properties
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.title:
                    metadata['title'] = core_props.title
            except:
                pass
            
            logger.info(
                f"✓ Loaded DOCX: {file_path.name} "
                f"({len(text)} chars, {metadata['paragraphs']} paragraphs)"
            )
            
            return {
                'text': text.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {str(e)}")
            raise
    
    @staticmethod
    def load_txt(file_path: Path) -> Dict[str, any]:
        """
        Load plain text file with encoding detection.
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Dict with 'text' and 'metadata' keys
        """
        try:
            # Try UTF-8 first, fallback to latin-1
            encodings = ['utf-8', 'latin-1', 'cp1252']
            text = None
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    encoding_used = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise ValueError(f"Could not decode file: {file_path.name}")
            
            metadata = {
                'source': str(file_path),
                'filename': file_path.name,
                'format': 'txt',
                'title': file_path.stem,
                'encoding': encoding_used,
                'lines': len(text.split('\n'))
            }
            
            logger.info(
                f"✓ Loaded TXT: {file_path.name} "
                f"({len(text)} chars, {metadata['lines']} lines)"
            )
            
            return {
                'text': text.strip(),
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {str(e)}")
            raise
    
    @classmethod
    def load(cls, file_path: str | Path) -> Dict[str, any]:
        """
        Auto-detect format and load document.
        
        Args:
            file_path: Path to document (str or Path)
            
        Returns:
            Dict with 'text' and 'metadata'
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If format not supported
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        suffix = file_path.suffix.lower()
        
        loaders = {
            '.pdf': cls.load_pdf,
            '.docx': cls.load_docx,
            '.doc': cls.load_docx,
            '.txt': cls.load_txt
        }
        
        if suffix not in loaders:
            raise ValueError(
                f"Unsupported format: {suffix}. "
                f"Supported: {', '.join(loaders.keys())}"
            )
        
        return loaders[suffix](file_path)


def load_document(file_path: str | Path) -> Dict[str, any]:
    """Convenience function for loading documents"""
    return DocumentLoader.load(file_path)
